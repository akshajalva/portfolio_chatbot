# Import necessary libraries and modules
# from langchain_ollama import OllamaLLM
import os
from dotenv import load_dotenv
from langchain_core.prompts import ChatPromptTemplate
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.chains import retrieval
from langchain.chains.combine_documents import create_stuff_documents_chain
from fastapi import FastAPI
from pydantic import BaseModel
from fastapi.middleware.cors import CORSMiddleware
from langchain.memory import ConversationBufferMemory
from docx import Document as DocxDocument
from langchain.schema import Document as LangchainDocument
from langchain_groq import ChatGroq


# Initialize FastAPI app
app = FastAPI()

# Load environment variables from a .env file
load_dotenv()
groq_api_key = os.getenv('GROQ_API_KEY')

# Initialize the Groq model using the provided API key
model = ChatGroq(groq_api_key=groq_api_key, model_name="llama-3.2-3b-preview")

# Enable CORS middleware to allow cross-origin requests
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Replace '*' with specific frontend URLs if you know them
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Function to read and chunk the document based on specific sections
def read_and_chunk_document(file_path):
    # Open and read the .docx file
    doc = DocxDocument(file_path)
    text = '\n'.join([para.text for para in doc.paragraphs])

    # Sections to split the document by
    sections = ["About Me", "Education", "Work Experience", "Skills", "Projects", "Website", "Personal Interests", "Career Gap"]
    # Split the document based on these sections
    chunks = []
    for i, section in enumerate(sections):
        start_idx = text.find(section)
        if start_idx != -1:
            end_idx = text.find(sections[i + 1], start_idx) if i + 1 < len(sections) else len(text)
            chunks.append(text[start_idx:end_idx].strip())
    return chunks

# Load and chunk the document
file_path = 'Aboutakshaj.docx'  # Replace with your .docx file path
chunked_documents = read_and_chunk_document(file_path)
# Create LangChain documents from chunks
documents = [LangchainDocument(page_content=chunk) for chunk in chunked_documents]

# Initialize the embedding model using HuggingFace embeddings
embedding_model = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

# Try to create the FAISS vector store from the documents using the embedding model
try:
    vector_store = FAISS.from_documents(documents, embedding=embedding_model)
    # print("Vector Store created successfully.")
except Exception as e:
    print(f"Error creating Vector Store: {e}")
    raise

# Create the Retrieval-Augmented Generation (RAG) Chain
# Initialize the retriever from the vector store
retriever = vector_store.as_retriever()

# Add conversation memory
memory = ConversationBufferMemory(memory_key="chat_history", return_messages=True)

# Define the prompt template for the chatbot
template = """
You are Akshaj Alva, and your goal is to provide accurate and concise responses strictly based on the knowledge base provided. 

Guidelines:
1. Respond only with information directly relevant to the user's query from the context below.
2. If the query is unrelated to the context, reply with: "I'm sorry, I can only answer questions about Akshaj Alva's professional background."
3. Keep responses concise and to the point unless explicitly asked for details.
4. Do not ask the user about their interests or knowledge. 
5. Always answer in the first person, as if you are Akshaj Alva.
6. Maintain a professional tone, reflecting your skills and knowledge.

Context:
{context}

Chat History:
{chat_history}


User's Question: {input}

Answer:
"""

# Create the prompt from the template
prompt = ChatPromptTemplate.from_template(template)

# Define the Retrieval QA Chain to process the documents and answer queries
document_chain = create_stuff_documents_chain(model, prompt)
qa_chain = retrieval.create_retrieval_chain(retriever, document_chain)


# Define a Pydantic Model for Input (used in FastAPI)
class ChatRequest(BaseModel):
    input: str

# Create the FastAPI endpoint for chatbot interaction
@app.post("/chat")
def chat_endpoint(request: ChatRequest):
    # Get user input from the API request
    user_input = request.input

    # Clear the memory before processing the new input
    memory.clear()
    # Retrieve the answer using the QA chain
    result = qa_chain.invoke({"input": user_input})
    
    # Return the response to the user
    return {"response": result.get('answer', "I don't have that information about Akshaj.")}

# Function to handle user input during the conversation (local testing)
def handle_conversation():
    print("Type 'exit' to end the conversation.")
    while True:
        user_input = input("You: ")
        if user_input.lower() == "exit":
            break

        # Debugging retrieval: Retrieve relevant documents for the user input
        retrieved_docs = retriever.get_relevant_documents(user_input)

        # Format the prompt with the retrieved context
        formatted_prompt = prompt.format(context="\n".join([doc.page_content for doc in retrieved_docs]), input=user_input)

        # Attempt to retrieve the answer from the QA chain
        try:
            result = qa_chain.invoke({"input": user_input})
            print("Bot Response:", result.get("answer", "I don't know."))
        except Exception as e:
            print(f"Error during model invocation: {e}")

# Run the conversation handling function if running locally
if __name__ == "__main__":
    handle_conversation()
